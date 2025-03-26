import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# List of lab URLs
urls = [
    "https://www.stevens.edu/center-for-complex-systems-and-enterprises",
    "https://www.stevens.edu/center-for-quantum-science-and-engineering",
    "https://www.stevens.edu/research/research-centers-and-labs/center-for-decision-technologies/center-for-decision-technologies",
    "https://www.stevens.edu/center-for-environmental-systems",
    "https://www.stevens.edu/center-for-healthcare-innovation",
    "https://www.stevens.edu/center-for-neuromechanics",
    "https://www.stevens.edu/craft",
    "https://www.stevens.edu/davidson-laboratory",
    "https://www.stevens.edu/school-business/hanlon-financial-systems-center",
    "https://www.stevens.edu/highly-filled-materials-institute",
    "https://www.stevens.edu/stevens-center-for-sustainability",
    "https://www.stevens.edu/stevens-institute-for-artificial-intelligence",
    "https://www.stevens.edu/new-jersey-center-for-microchemical-systems"
]

# Initialize data storage
data = {
    "Research Center/Lab": [],
    "Description": [],
    "Details": []
}

# Function to scrape a single page
def scrape_lab_page(url):
    try:
        response = requests.get(url)
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Extract lab name (from title or h1)
        title = soup.find('h1') or soup.find('title')
        lab_name = title.text.strip() if title else url.split('/')[-1].replace('-', ' ').title()
        
        # Extract description (look for common patterns: p, div with class 'content', etc.)
        desc = ""
        for tag in ['p', 'div']:
            content = soup.find(tag, class_=['content', 'description', 'intro', None])
            if content and len(content.get_text(strip=True)) > 50:  # Ensure it's substantial
                desc = content.get_text(strip=True)
                break
        
        # Extract additional details (research focus, location, etc.)
        details = ""
        for section in soup.find_all(['h2', 'h3', 'section']):
            text = section.get_text(strip=True)
            if any(keyword in text.lower() for keyword in ['research', 'focus', 'mission', 'location', 'director']):
                details += f"{text}: {section.find_next('p').get_text(strip=True) if section.find_next('p') else ''}\n"
        
        return lab_name, desc or "No description found on page.", details or "No additional details found."
    except Exception as e:
        return url.split('/')[-1].replace('-', ' ').title(), f"Error scraping page: {str(e)}", ""

# Scrape all pages
for url in urls:
    lab_name, desc, details = scrape_lab_page(url)
    data["Research Center/Lab"].append(lab_name)
    data["Description"].append(desc)
    data["Details"].append(details)

# Create DataFrame
df = pd.DataFrame(data)

# Create Word document
doc = Document()
doc.add_heading('Stevens Research Network', 0)

# Add table
doc.add_heading('Research Centers and Labs - Overview Table', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Research Center/Lab'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'Details'

for _, row in df.iterrows():
    row_cells = table.add_row().cells
    row_cells[0].text = row['Research Center/Lab']
    row_cells[1].text = row['Description']
    row_cells[2].text = row['Details']

# Add detailed passages
doc.add_heading('Detailed Information', level=1)
for _, row in df.iterrows():
    doc.add_heading(row['Research Center/Lab'], level=2)
    doc.add_paragraph(f"Description: {row['Description']}")
    doc.add_paragraph(f"Additional Details:\n{row['Details']}")

# Save document
doc_filename = 'stevens_research_network.docx'
doc.save(doc_filename)
print(f"Document saved as '{doc_filename}'")